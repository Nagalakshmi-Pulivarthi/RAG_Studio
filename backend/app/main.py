import json
import queue
import threading

from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from pathlib import Path
from io import BytesIO

from app.config import ANTHROPIC_API_KEY


app = FastAPI(title="RAG API")


@app.on_event("startup")
def _ensure_rag_deps():
    """Fail fast if faiss/sentence-transformers are missing (e.g. wrong Python when using --reload on Windows)."""
    try:
        import faiss  # noqa: F401
        from sentence_transformers import SentenceTransformer  # noqa: F401
    except ModuleNotFoundError as e:
        raise RuntimeError(
            "RAG dependencies not found. Start the server with the venv Python:\n"
            "  cd backend\n"
            "  .\\venv\\Scripts\\python.exe run.py\n"
            f"Missing: {e.name}"
        ) from e


@app.exception_handler(Exception)
def unhandled_exception_handler(request, exc: Exception):
    """Return the real error as JSON so we can debug 500s."""
    if isinstance(exc, HTTPException):
        return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})
    return JSONResponse(
        status_code=500,
        content={"detail": f"Error: {exc!s}", "type": type(exc).__name__},
    )

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class IngestBody(BaseModel):
    text: str | None = None
    path: str | None = None


class IngestUrlBody(BaseModel):
    url: str


class ChatBody(BaseModel):
    message: str
    session_id: str | None = None


def _is_numeric_cell(s: str) -> bool:
    """Return True if s looks like a number or quantity (e.g. 680, 43g, 860mg, 2%)."""
    import re
    return bool(re.match(r"^[\d,]+\.?\d*\s*(g|mg|cal|kcal|%|oz|ml|lb|mcg|iu)?$", s.strip(), re.IGNORECASE))


def _looks_like_header(row: list[str]) -> bool:
    """Return True when a row is mostly text labels rather than numbers — i.e. a header row."""
    non_empty = [c for c in row if c]
    if not non_empty:
        return False
    numeric_count = sum(1 for c in non_empty if _is_numeric_cell(c))
    return numeric_count / len(non_empty) < 0.4


def _table_to_sentences(table: list[list]) -> list[str]:
    """
    Convert a pdfplumber table (list of rows) into natural-language sentences so that
    the embedding model can link item names to their values.

    Input:  [["Item", "Calories", "Fat"], ["Chuck Wagon (Plain)", "680", "43g"], ...]
    Output: ["Chuck Wagon (Plain): 680 Calories, 43g Fat."]
    """
    if not table:
        return []

    # Normalise cells to stripped strings
    cleaned = [[str(c or "").strip() for c in row] for row in table]

    # Detect header
    if len(cleaned) >= 2 and _looks_like_header(cleaned[0]):
        headers = cleaned[0]
        data_rows = cleaned[1:]
    else:
        headers = []
        data_rows = cleaned

    sentences: list[str] = []
    for row in data_rows:
        non_empty = [c for c in row if c]
        if not non_empty:
            continue

        if headers and row and row[0]:
            item_name = row[0]
            pairs = []
            for i in range(1, len(row)):
                val = row[i] if i < len(row) else ""
                col = headers[i] if i < len(headers) else ""
                if val and col:
                    pairs.append(f"{val} {col}")
                elif val:
                    pairs.append(val)
            if pairs:
                sentences.append(f"{item_name}: {', '.join(pairs)}.")
            else:
                sentences.append(item_name)
        else:
            sentences.append(" | ".join(non_empty))

    return sentences


def _extract_pdf_text(data: bytes) -> str:
    """
    Extract text from a PDF byte stream.

    Uses pdfplumber for table-aware extraction: each table row is converted to a
    natural-language sentence so the embedding model can match queries to specific
    items (e.g. 'Chuck Wagon (Plain): 680 Calories, 43g Fat, 860mg Sodium.').
    Non-table text is extracted separately to avoid duplication.
    Falls back to pypdf if pdfplumber is unavailable or fails.
    """
    try:
        import pdfplumber
    except ImportError:
        pdfplumber = None  # type: ignore

    if pdfplumber is not None:
        try:
            parts: list[str] = []
            with pdfplumber.open(BytesIO(data)) as pdf:
                for page in pdf.pages:
                    tbl_objects = page.find_tables()

                    if tbl_objects:
                        # Convert each table to natural language sentences
                        table_bboxes = []
                        for tbl_obj in tbl_objects:
                            table_bboxes.append(tbl_obj.bbox)
                            raw_table = tbl_obj.extract()
                            for sentence in _table_to_sentences(raw_table):
                                parts.append(sentence)

                        # Extract text that falls outside every table bounding box
                        words = page.extract_words()
                        outside = [
                            w["text"] for w in words
                            if not any(
                                w["x0"] >= bbox[0] and w["top"] >= bbox[1]
                                and w["x1"] <= bbox[2] and w["bottom"] <= bbox[3]
                                for bbox in table_bboxes
                            )
                        ]
                        if outside:
                            parts.append(" ".join(outside))
                    else:
                        txt = page.extract_text() or ""
                        if txt.strip():
                            parts.append(txt)

            return "\n\n".join(p for p in parts if p.strip())
        except Exception:
            pass  # fall through to pypdf fallback

    # pypdf fallback
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(data))
    parts = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt:
            parts.append(txt)
    return "\n".join(parts)


def _extract_html_text(html: str) -> str:
    """Extract visible text from HTML."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator="\n", strip=True)


def _extract_docx_text(data: bytes) -> str:
    """Extract text from a Word .docx file."""
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/")
def root():
    """Root redirect/info so we can confirm this is the RAG backend."""
    return {"message": "RAG API", "docs": "/docs", "health": "/health"}


@app.get("/api/info")
def api_info():
    """List main API endpoints (helps verify this is the RAG backend)."""
    return {
        "message": "RAG API",
        "endpoints": [
            "POST /api/ingest",
            "POST /api/ingest-upload",
            "POST /api/ingest-url",
            "POST /api/clear-index",
            "POST /api/chat",
            "GET /api/check-deps",
            "GET /api/index-status",
        ],
    }


@app.get("/api/index-status")
def api_index_status():
    """Return the unique sources and total chunk count currently in the index."""
    from app.ingest import get_index_and_docstore
    from collections import Counter

    _, docstore = get_index_and_docstore()
    if not docstore:
        return {"total_chunks": 0, "sources": []}

    counts = Counter(doc.get("source", "unknown") for doc in docstore)
    sources = [{"source": s, "chunks": c} for s, c in counts.items()]
    return {"total_chunks": len(docstore), "sources": sources}


@app.get("/api/check-deps")
def check_deps():
    """Returns 200 if faiss and sentence_transformers are importable (this process has the right Python/venv)."""
    try:
        import faiss  # noqa: F401
        from sentence_transformers import SentenceTransformer  # noqa: F401
        return {"ok": True, "message": "RAG dependencies (faiss, sentence_transformers) are available."}
    except ModuleNotFoundError as e:
        raise HTTPException(
            status_code=503,
            detail=(
                f"Missing module: {e.name}. This server process is not using the venv. "
                "Stop this server, then start it with: cd backend && .\\venv\\Scripts\\python.exe -m uvicorn app.main:app (no --reload)."
            ),
        )


@app.post("/api/ingest")
def api_ingest(body: IngestBody):
    """Ingest text or a file. Send either 'text' or 'path' in the body (not both)."""
    try:
        from app.ingest import ingest_text, ingest_file
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Could not load ingest module: {e!s}")

    # Prefer text when the client sent it, so pasted text is never confused with a path.
    if body.text is not None:
        text = (body.text or "").strip()
        if not text:
            raise HTTPException(
                status_code=400,
                detail="Text cannot be empty. Paste your content in the text box and try again.",
            )
        try:
            result = ingest_text(text, source="pasted")
            return result
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Ingest failed: {e!s}")
    if body.path and body.path.strip():
        try:
            result = ingest_file(body.path.strip())
            return result
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="File not found")
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Ingest failed: {e!s}")
    raise HTTPException(
        status_code=400,
        detail="Provide either 'text' or 'path' in the request body",
    )


@app.get("/api/ingest-upload")
def api_ingest_upload_get():
    """Upload endpoint only accepts POST. Confirms this is the RAG backend if you get 405 here."""
    raise HTTPException(status_code=405, detail="Use POST with multipart/form-data and a 'file' field")


@app.post("/api/ingest-upload")
async def api_ingest_upload(file: UploadFile = File(...)):
    """
    Ingest an uploaded PDF, Word (.docx), or HTML file.

    - For PDFs: extract text from all pages.
    - For .docx: extract paragraphs and table text.
    - For HTML: strip tags and boilerplate, keeping visible text.
    """
    try:
        from app.ingest import ingest_text
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Could not load ingest module: {e!s}")

    raw = await file.read()
    filename = file.filename or "upload"
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    try:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            text = _extract_pdf_text(raw)
        elif lower.endswith(".docx"):
            text = _extract_docx_text(raw)
        else:
            try:
                html = raw.decode("utf-8", errors="ignore")
            except Exception:
                html = raw.decode("latin1", errors="ignore")
            text = _extract_html_text(html)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not extract text from file: {e!s}")

    if not text.strip():
        return {"chunks_added": 0, "message": "No text extracted from file", "source": filename}

    result = ingest_text(text, source=filename)
    return result


@app.post("/api/ingest-url")
async def api_ingest_url(body: IngestUrlBody):
    """
    Fetch a URL (e.g. a web page), extract text from the HTML, and ingest it into the RAG index.
    """
    import httpx

    url = (body.url or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="URL is required")
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL must start with http:// or https://")

    try:
        from app.ingest import ingest_text
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Could not load ingest module: {e!s}")

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
            response = await client.get(url)
            response.raise_for_status()
            raw = response.text
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="Request timed out. The URL took too long to respond.")
    except httpx.HTTPStatusError as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to fetch URL: {e.response.status_code} {e.response.reason_phrase}",
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not fetch URL: {e!s}")

    text = _extract_html_text(raw)
    if not text.strip():
        return {"chunks_added": 0, "message": "No text extracted from page", "source": url}

    source_label = url if len(url) <= 80 else url[:77] + "..."
    result = ingest_text(text, source=source_label)
    return result


@app.post("/api/clear-index")
def api_clear_index():
    """Clear the RAG index and docstore. Next ingest will start fresh (useful when you want only pasted/uploaded content)."""
    try:
        from app.ingest import clear_index
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Could not load ingest module: {e!s}")
    return clear_index()


@app.post("/api/chat")
def api_chat(body: ChatBody):
    """Answer using the Agno agent (search_documents tool + Claude)."""
    try:
        from app.agent import get_agent
        from app import tools as _tools
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Could not load agent module: {e!s}")

    if not ANTHROPIC_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="ANTHROPIC_API_KEY is not set. Add it to .env (get a key at https://console.anthropic.com).",
        )
    message = (body.message or "").strip()
    if not message:
        raise HTTPException(status_code=400, detail="message is required")
    try:
        agent = get_agent()
        run_kwargs = {"session_id": body.session_id} if body.session_id else {}
        response = agent.run(message, **run_kwargs)
        answer = response.content or ""
        sources = [{"source": c["source"], "text": c["text"][:200]} for c in _tools._last_chunks]
        return {"answer": answer, "sources": sources}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Chat failed: {e!s}")


@app.post("/api/chat/stream")
def api_chat_stream(body: ChatBody):
    """Stream agent events as SSE: tool_call → tool_done → answer."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=503, detail="ANTHROPIC_API_KEY is not set.")
    message = (body.message or "").strip()
    if not message:
        raise HTTPException(status_code=400, detail="message is required")

    try:
        from app.agent import get_agent
        from app import tools as _tools
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Could not load agent module: {e!s}")

    event_queue: queue.Queue = queue.Queue()

    session_id = body.session_id

    def run_agent():
        _tools._event_queue = event_queue
        try:
            agent = get_agent()
            run_kwargs = {"session_id": session_id} if session_id else {}
            response = agent.run(message, **run_kwargs)
            sources = [{"source": c["source"], "text": c["text"][:200]} for c in _tools._last_chunks]
            event_queue.put({"type": "answer", "content": response.content or "", "sources": sources})
        except Exception as exc:
            event_queue.put({"type": "error", "message": str(exc)})
        finally:
            _tools._event_queue = None
            event_queue.put(None)  # sentinel — signals generator to stop

    def generate():
        thread = threading.Thread(target=run_agent, daemon=True)
        thread.start()
        try:
            while True:
                try:
                    event = event_queue.get(timeout=120)
                except queue.Empty:
                    yield f"data: {json.dumps({'type': 'error', 'message': 'Agent timed out'})}\n\n"
                    break
                if event is None:
                    break
                yield f"data: {json.dumps(event)}\n\n"
        finally:
            thread.join(timeout=5)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


